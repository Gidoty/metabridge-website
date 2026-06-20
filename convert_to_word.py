#!/usr/bin/env python3
"""Convert MetaBridge markdown curriculum files to formatted Word documents.

Replicates the original MetaBridge document visual style:
  - White background with clean typography
  - Navy section banners (#1B3A6B) for ## headings
  - Dark green sub-section headings (#2E7D32) for ### headings
  - Callout boxes with coloured left borders (objectives, labs, key concepts)
  - Navy table headers with white text
  - Page header and footer with MetaBridge branding
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette (from original MetaBridge documents) ─────────────────────
NAVY_BANNER  = RGBColor(0x1B, 0x3A, 0x6B)  # section banners, table headers
NAVY_DARK    = RGBColor(0x0C, 0x1A, 0x2E)  # very dark navy (H1)
GREEN_SUB    = RGBColor(0x2E, 0x7D, 0x32)  # subsection headings (###)
GOLD_CONCEPT = RGBColor(0xC8, 0xA0, 0x00)  # key concept headers
BLUE_OBJ     = RGBColor(0x0D, 0x49, 0x99)  # module objective border
BLACK        = RGBColor(0x1A, 0x1A, 0x1A)  # body text
MUTED        = RGBColor(0x64, 0x74, 0x8B)  # muted / footer text
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
GREY_HEADER  = RGBColor(0x90, 0x9A, 0xA8)  # page header text

# Course accent colours (for cover page and H1)
COURSE_COLOURS = {
    "cybersecurity":             RGBColor(0xC0, 0x39, 0x2B),  # dark red
    "data-analytics":            RGBColor(0x00, 0x7A, 0x99),  # dark cyan
    "ai-prompt-engineering":     RGBColor(0x6C, 0x3A, 0x9E),  # dark purple
    "blockchain-cryptocurrency": RGBColor(0xB8, 0x76, 0x00),  # dark gold
}

COURSE_LABELS = {
    "cybersecurity":             "Cybersecurity",
    "data-analytics":            "Data Analytics",
    "ai-prompt-engineering":     "AI & Prompt Engineering",
    "blockchain-cryptocurrency": "Blockchain & Cryptocurrency",
}

DOC_LABELS = {
    "01-curriculum":       "Curriculum Overview",
    "01-curriculum-part2": "Curriculum Overview — Part 2",
    "02-lesson-notes":     "Lesson Notes",
    "04-kahoot-quiz-bank": "Kahoot! Quiz Bank",
}


# ── XML / low-level helpers ──────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour):
    """Set a table cell's background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


def _set_para_shading(paragraph, hex_fill):
    """Set a paragraph's background shading colour."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    # Remove existing shd first
    for existing in pPr.findall(qn('w:shd')):
        pPr.remove(existing)
    pPr.append(shd)


def _set_para_left_border(paragraph, hex_colour, size=36, space=8):
    """Add a thick left border to a paragraph (callout box style)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), str(space))
    left.set(qn('w:color'), hex_colour)
    pBdr.append(left)
    # Remove existing pBdr first
    for existing in pPr.findall(qn('w:pBdr')):
        pPr.remove(existing)
    pPr.append(pBdr)


def _add_page_break(doc):
    from docx.oxml.ns import qn as _qn
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(_qn('w:type'), 'page')
    run._r.append(br)
    return p


def _add_para_spacing(paragraph, before=0, after=6, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


# ── Header / footer ──────────────────────────────────────────────────────────

def _add_header_footer(doc, course_name, doc_type, accent):
    """Add running page header and footer."""
    section = doc.sections[0]

    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"MetaBridge Academy  ·  {course_name}  ·  2025")
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY_HEADER
    run.font.name = 'Calibri'
    run.font.italic = True

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Left side: confidentiality notice
    r1 = fp.add_run("Confidential — MetaBridge Academy Internal Document    Page ")
    r1.font.size = Pt(8)
    r1.font.color.rgb = GREY_HEADER
    r1.font.name = 'Calibri'

    # Inline page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r1._r.append(fldChar1)
    r1._r.append(instrText)
    r1._r.append(fldChar2)


# ── Document styles ──────────────────────────────────────────────────────────

def _setup_styles(doc, accent):
    s = doc.styles

    # Normal
    n = s['Normal']
    n.font.name = 'Calibri'
    n.font.size = Pt(11)
    n.font.color.rgb = BLACK
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = Pt(15)

    # H1 — document title (accent colour, large)
    h1 = s['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = accent
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(8)

    # H2 — module / major section (will be formatted as navy banner)
    h2 = s['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = WHITE
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(4)

    # H3 — sub-section (dark green)
    h3 = s['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = GREEN_SUB
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # H4 — sub-sub-section (dark navy)
    h4 = s['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.color.rgb = NAVY_DARK
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(2)

    # List Bullet
    try:
        lb = s['List Bullet']
    except Exception:
        lb = s.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    lb.font.name = 'Calibri'
    lb.font.size = Pt(11)
    lb.font.color.rgb = BLACK
    lb.paragraph_format.left_indent = Inches(0.3)
    lb.paragraph_format.space_after = Pt(3)

    # List Bullet 2 (nested)
    try:
        lb2 = s['List Bullet 2']
    except Exception:
        lb2 = s.add_style('List Bullet 2', WD_STYLE_TYPE.PARAGRAPH)
    lb2.font.name = 'Calibri'
    lb2.font.size = Pt(10.5)
    lb2.font.color.rgb = BLACK
    lb2.paragraph_format.left_indent = Inches(0.6)
    lb2.paragraph_format.space_after = Pt(2)

    # List Number
    try:
        ln = s['List Number']
    except Exception:
        ln = s.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
    ln.font.name = 'Calibri'
    ln.font.size = Pt(11)
    ln.font.color.rgb = BLACK
    ln.paragraph_format.left_indent = Inches(0.3)
    ln.paragraph_format.space_after = Pt(3)


# ── Cover page ───────────────────────────────────────────────────────────────

def _add_cover_page(doc, course_name, doc_type, accent):
    """Styled cover page with navy banner."""
    # Top navy banner paragraph
    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_shading(p_banner, '1B3A6B')
    _add_para_spacing(p_banner, before=0, after=0)

    r = p_banner.add_run("  MetaBridge Academy  ")
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = WHITE

    # Sub-banner with course name
    p_course = doc.add_paragraph()
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_shading(p_course, '243A5E')
    _add_para_spacing(p_course, before=0, after=0)

    r2 = p_course.add_run(f"  {course_name.upper()}  ")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0xB0, 0xC8, 0xE8)

    # Spacer
    sp = doc.add_paragraph()
    _add_para_spacing(sp, before=4, after=4)

    # Document type (large, accent coloured)
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_spacing(p_type, before=12, after=4)
    r3 = p_type.add_run(doc_type.upper())
    r3.font.name = 'Calibri'
    r3.font.size = Pt(18)
    r3.font.bold = True
    r3.font.color.rgb = accent

    # Year / edition
    p_yr = doc.add_paragraph()
    p_yr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_spacing(p_yr, before=0, after=24)
    r4 = p_yr.add_run("2025 Edition  ·  metabridgeacademy.com")
    r4.font.name = 'Calibri'
    r4.font.size = Pt(10)
    r4.font.color.rgb = MUTED

    # Thin bottom rule
    _add_horizontal_rule(doc, color_hex='1B3A6B')
    doc.add_paragraph()


# ── Callout box helpers ──────────────────────────────────────────────────────

def _add_callout_box(doc, header_text, body_lines, bg_hex, border_hex, header_colour):
    """Styled callout box: coloured bg + thick left border."""
    # Header paragraph
    p_hdr = doc.add_paragraph()
    _set_para_shading(p_hdr, bg_hex)
    _set_para_left_border(p_hdr, border_hex, size=48, space=6)
    _add_para_spacing(p_hdr, before=6, after=0)
    p_hdr.paragraph_format.left_indent = Inches(0.15)

    r = p_hdr.add_run(header_text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = header_colour

    # Body paragraphs
    for bline in body_lines:
        pb = doc.add_paragraph()
        _set_para_shading(pb, bg_hex)
        _set_para_left_border(pb, border_hex, size=48, space=6)
        _add_para_spacing(pb, before=0, after=1)
        pb.paragraph_format.left_indent = Inches(0.15)
        _apply_inline(pb, bline)

    # Closing spacer
    sp = doc.add_paragraph()
    _add_para_spacing(sp, before=0, after=6)


def _add_horizontal_rule(doc, color_hex='1B3A6B'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    _add_para_spacing(p, before=0, after=2)
    return p


# ── Inline formatting helper ─────────────────────────────────────────────────

def _apply_inline(paragraph, text, default_colour=None):
    """Parse **bold**, `code`, and plain text runs."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            if default_colour:
                run.font.color.rgb = default_colour
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x0F, 0x77, 0x2A)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            if default_colour:
                run.font.color.rgb = default_colour


# ── Table parser ─────────────────────────────────────────────────────────────

def _parse_table(doc, lines, start_idx, accent):
    """Parse a markdown table. Returns next line index."""
    table_lines = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        table_lines.append(lines[i])
        i += 1

    if len(table_lines) < 2:
        return i

    data_rows = [l for l in table_lines if not re.match(r'^\|[-:\s|]+\|$', l.strip())]
    if not data_rows:
        return i

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    if not parsed:
        return i

    col_count = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = 'Table Grid'

    # Set column widths proportionally
    available = Inches(6.2)
    col_w = int(available / col_count)
    for col in table.columns:
        col.width = col_w

    for ri, row in enumerate(parsed):
        for ci in range(col_count):
            cell_text = row[ci] if ci < len(row) else ''
            # Strip markdown bold/code from cell text
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            cell = table.cell(ri, ci)
            cell.text = clean

            p = cell.paragraphs[0]
            if ri == 0:
                # Header row: navy background, white bold text
                _set_cell_bg(cell, '1B3A6B')
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = WHITE
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
            else:
                # Alternating light grey for even rows
                if ri % 2 == 0:
                    _set_cell_bg(cell, 'F5F7FA')
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = BLACK

    doc.add_paragraph()
    return i


# ── State machine for callout detection ─────────────────────────────────────

class _BlockState:
    NORMAL = 'normal'
    BLOCKQUOTE = 'blockquote'
    CODE = 'code'


def _detect_callout_type(header_text):
    """Detect which type of callout box a blockquote should be."""
    t = header_text.upper()
    if 'OBJECTIVE' in t or 'OUTCOME' in t or 'BY THE END' in t:
        return 'objective'  # blue
    if 'LAB' in t or 'PRACTICAL' in t or 'EXERCISE' in t or 'HANDS-ON' in t:
        return 'lab'        # green
    if any(kw in t for kw in ['KEY CONCEPT', 'CIA TRIAD', 'SOCIAL ENGINEERING',
                                'IMPORTANT', 'NOTE:', 'REMEMBER', 'QUIZ FOCUS',
                                'WARNING', 'CRITICAL']):
        return 'concept'    # gold/yellow
    if 'LEVEL:' in t or 'ESTIMATED' in t or 'DURATION' in t:
        return 'info'       # light blue
    return 'generic'        # default blue


# ── Main converter ───────────────────────────────────────────────────────────

def convert_md_to_docx(md_path, docx_path, course_key):
    accent = COURSE_COLOURS.get(course_key, RGBColor(0x1B, 0x3A, 0x6B))
    course_name = COURSE_LABELS.get(course_key, course_key)

    stem = os.path.splitext(os.path.basename(md_path))[0]
    doc_type = DOC_LABELS.get(stem, stem)

    doc = Document()

    # Page margins (generous like original)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    _setup_styles(doc, accent)
    _add_header_footer(doc, course_name, doc_type, accent)
    _add_cover_page(doc, course_name, doc_type, accent)

    with open(md_path, 'r', encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')
    i = 0
    state = _BlockState.NORMAL
    code_lines = []
    code_lang = ''
    bq_lines = []   # accumulated blockquote lines

    def flush_blockquote():
        nonlocal bq_lines
        if not bq_lines:
            return
        first_line = bq_lines[0] if bq_lines else ''
        callout = _detect_callout_type(first_line)

        if callout == 'lab':
            bg, border, hc = 'E8F5E9', '2E7D32', GREEN_SUB
        elif callout == 'concept':
            bg, border, hc = 'FFF9E6', 'C8A000', GOLD_CONCEPT
        elif callout in ('objective', 'info', 'generic'):
            bg, border, hc = 'EBF0F8', '0D4999', BLUE_OBJ
        else:
            bg, border, hc = 'EBF0F8', '0D4999', BLUE_OBJ

        _add_callout_box(doc, first_line, bq_lines[1:], bg, border, hc)
        bq_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code block ──────────────────────────────────────────────────────
        if stripped.startswith('```'):
            if state != _BlockState.CODE:
                flush_blockquote()
                state = _BlockState.CODE
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                # End code block
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x0F, 0x77, 0x2A)
                    p.paragraph_format.left_indent = Inches(0.3)
                    _add_para_spacing(p, before=4, after=4)
                    _set_para_shading(p, 'F4F6F9')
                state = _BlockState.NORMAL
                code_lines = []
            i += 1
            continue

        if state == _BlockState.CODE:
            code_lines.append(line)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────────────
        if stripped.startswith('> '):
            bq_content = stripped[2:]
            # Strip markdown bold markers for display
            bq_content = re.sub(r'\*\*([^*]+)\*\*', r'\1', bq_content)
            bq_lines.append(bq_content)
            i += 1
            continue
        else:
            # Any non-blockquote line flushes pending blockquote
            flush_blockquote()

        # ── Skip empty lines ─────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped) or stripped in ('***', '___'):
            _add_horizontal_rule(doc, '1B3A6B')
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        if stripped.startswith('#### '):
            p = doc.add_paragraph(style='Heading 4')
            _apply_inline(p, stripped[5:])
            _add_para_spacing(p, before=8, after=2)
            i += 1
            continue

        if stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_paragraph(style='Heading 3')
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.color.rgb = GREEN_SUB
            _add_para_spacing(p, before=12, after=3)
            # Thin green underline rule
            _add_horizontal_rule(doc, '2E7D32')
            i += 1
            continue

        if stripped.startswith('## '):
            text = stripped[3:]
            p = doc.add_paragraph(style='Heading 2')
            _set_para_shading(p, '1B3A6B')
            _add_para_spacing(p, before=18, after=4)
            p.paragraph_format.left_indent = Inches(0.1)
            run = p.add_run(f"  {text}")
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = WHITE
            i += 1
            continue

        if stripped.startswith('# '):
            text = stripped[2:]
            p = doc.add_paragraph(style='Heading 1')
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = accent
            _add_para_spacing(p, before=0, after=8)
            i += 1
            continue

        # ── Markdown table ────────────────────────────────────────────────────
        if (stripped.startswith('|')
                and i + 1 < len(lines)
                and re.match(r'^\|[-:\s|]+\|$', lines[i + 1].strip())):
            i = _parse_table(doc, lines, i, accent)
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            _apply_inline(p, m.group(2))
            _add_para_spacing(p, before=0, after=3)
            i += 1
            continue

        # ── Nested bullet (2+ spaces) ─────────────────────────────────────────
        if re.match(r'^\s{2,}[-*]\s+', line):
            p = doc.add_paragraph(style='List Bullet 2')
            text = re.sub(r'^\s*[-*]\s+', '', line).strip()
            _apply_inline(p, text)
            _add_para_spacing(p, before=0, after=2)
            i += 1
            continue

        # ── Bullet ────────────────────────────────────────────────────────────
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _apply_inline(p, stripped[2:])
            _add_para_spacing(p, before=0, after=3)
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        p = doc.add_paragraph()
        _apply_inline(p, stripped)
        _add_para_spacing(p, before=0, after=6)
        i += 1

    # Flush any remaining blockquote
    flush_blockquote()

    doc.save(docx_path)
    print(f"  Saved: {docx_path}")


# ── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    base = '/home/user/metabridge-website/curriculum'
    out_dir = '/home/user/metabridge-website/exports'
    os.makedirs(out_dir, exist_ok=True)

    files = [
        ('cybersecurity', '01-curriculum.md'),
        ('cybersecurity', '01-curriculum-part2.md'),
        ('cybersecurity', '02-lesson-notes.md'),
        ('cybersecurity', '04-kahoot-quiz-bank.md'),
        ('data-analytics', '01-curriculum.md'),
        ('data-analytics', '01-curriculum-part2.md'),
        ('data-analytics', '02-lesson-notes.md'),
        ('data-analytics', '04-kahoot-quiz-bank.md'),
        ('ai-prompt-engineering', '01-curriculum.md'),
        ('ai-prompt-engineering', '01-curriculum-part2.md'),
        ('ai-prompt-engineering', '02-lesson-notes.md'),
        ('ai-prompt-engineering', '04-kahoot-quiz-bank.md'),
        ('blockchain-cryptocurrency', '01-curriculum.md'),
        ('blockchain-cryptocurrency', '01-curriculum-part2.md'),
        ('blockchain-cryptocurrency', '02-lesson-notes.md'),
        ('blockchain-cryptocurrency', '04-kahoot-quiz-bank.md'),
    ]

    for course, filename in files:
        stem = filename.replace('.md', '')
        label = DOC_LABELS.get(stem, stem).replace(' ', '-').replace('/', '-').replace('—', '')
        course_label = COURSE_LABELS.get(course, course).replace(' ', '-').replace('&', 'and')
        out_name = f"{course_label}-{label}.docx"
        src = os.path.join(base, course, filename)
        dst = os.path.join(out_dir, out_name)
        if os.path.exists(src):
            print(f"Converting: {course}/{filename}")
            convert_md_to_docx(src, dst, course)
        else:
            print(f"  MISSING: {src}")

    print("\nWord conversion complete.")
